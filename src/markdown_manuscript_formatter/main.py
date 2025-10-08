import yaml
import os
import argparse
import re
import smartypants
import html
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn

# --- Configuration ---
# The configuration file and output directory are now handled by command-line arguments.
def add_smart_text(paragraph, text):
    """
    Processes text with smartypants and simple markdown, adding it to a 
    docx paragraph run by run.
    """
    # Split text by our markdown tokens, keeping the tokens.
    # This allows us to process text segment by segment.
    # Tokens: ***, **, *, _
    tokens = re.split(r'(\*{1,3}|_{1})', text)

    is_bold = False
    is_italic = False

    for token in tokens:
        if token == '***':
            is_bold = not is_bold
            is_italic = not is_italic
        elif token == '**':
            is_bold = not is_bold
        elif token == '*' or token == '_':
            is_italic = not is_italic
        elif token:
            # Process with smartypants *after* splitting by markdown.
            # This prevents markdown characters from confusing smartypants.
            processed_token = html.unescape(smartypants.smartypants(token))
            run = paragraph.add_run(processed_token)
            run.bold = is_bold
            run.italic = is_italic

def create_title_page(config, doc, story_type):
    """Generates a standard manuscript title page."""
    author = config.get('author', {})
    metadata = config.get('metadata', {})
    
    # Use a 1x2 table to align contact info (left) and word count (right).
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False # Allow manual column sizing
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)

    # --- Left Cell: Author Contact Info ---
    left_cell = table.cell(0, 0)
    p_contact = left_cell.paragraphs[0] # Use the cell's default paragraph
    
    contact_parts = [
        author.get('legal_name'),
        author.get('street_address'),
        author.get('city_state_zip'),
        author.get('phone'),
        author.get('email')
    ]
    # Filter out any None values and join with newlines.
    # add_run correctly handles '\n' as a line break within a paragraph.
    contact_block = "\n".join(part for part in contact_parts if part)
    add_smart_text(p_contact, contact_block)

    # --- Right Cell: Word Count ---
    right_cell = table.cell(0, 1)
    p_word_count = right_cell.paragraphs[0]
    p_word_count.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Add a placeholder. We'll replace this after counting all words.
    p_word_count.add_run("Approx. _WORD_COUNT_PLACEHOLDER_ words")

    # --- Make Table Borders Invisible ---
    # This is a robust, low-level method to "hide" a table for layout.
    # It avoids relying on style names which may not exist in the template.
    no_border_xml = (
        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/>'
        r'<w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
        r'</w:tcBorders>'
    )
    
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_pr.append(parse_xml(no_border_xml))

    # Center the title and byline
    title = metadata.get('title', 'Untitled Novel').upper()
    byline = f"by\n{metadata.get('byline', 'Anonymous')}"
    
    # Add empty paragraphs for spacing
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(title)
    run_title.bold = True
    p_byline = doc.add_paragraph(byline)
    p_byline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if story_type == 'novel':
        doc.add_page_break()
    else: # For short stories, add spacing before the text starts on the same page.
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

def setup_header(doc, config):
    """
    Adds a right-aligned header to all pages except the first.
    Header format: LastName | Short Title | PageNumber
    """
    metadata = config.get('metadata', {})
    last_name = metadata.get('last_name', 'Author')
    short_title = metadata.get('short_title', 'Manuscript')

    # Enable "Different first page" header/footer
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Access the primary header (for pages > 1)
    header = section.header
    p_header = header.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add header text in multiple runs to apply formatting.
    p_header.add_run(f"{last_name} | ")
    run_title = p_header.add_run(short_title)
    run_title.italic = True
    p_header.add_run(" | ")

    # Add the dynamic page number field
    # This requires manipulating the underlying OOXML.
    run = p_header.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run._r.append(instrText)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

def append_file_content(filepath, doc):
    """
    Reads markdown content from a given filepath, ignores heading lines,
    and appends it to the document.
    """
    try:
        with open(filepath, 'r', encoding='utf-8') as f_in:
            content = f_in.read()
            # Split content into paragraphs
            paragraphs = content.split('\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue # Skip empty lines between paragraphs
                if re.match(r'^#{1,6}\s', para_text):
                    continue # Skip markdown headings
                if para_text.startswith('%'):
                    continue # Skip comment lines

                p = doc.add_paragraph()
                # Set standard manuscript formatting
                p.paragraph_format.first_line_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 2.0 # Double-spaced
                add_smart_text(p, para_text)

    except FileNotFoundError:
        print(f"--> WARNING: Could not find file: {filepath}. It will be skipped.")
    except Exception as e:
        print(f"--> ERROR: An error occurred reading {filepath}: {e}")

def process_chapter(chapter_data, doc, story_type):
    """Processes a chapter dictionary and writes its content to the document."""
    heading = []
    if 'number' in chapter_data:
        heading.append(f"Chapter {chapter_data['number']}")
    if 'title' in chapter_data:
        heading.append(chapter_data['title'])
    
    if heading:
        full_heading = smartypants.smartypants(": ".join(heading))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(full_heading)
        run.bold = True
        # Add space after heading
        doc.add_paragraph()

    if 'file' in chapter_data:
        append_file_content(chapter_data['file'], doc)
    elif 'files' in chapter_data:
        file_paths = chapter_data.get('files', [])
        for i, file_path in enumerate(file_paths):
            append_file_content(file_path, doc)
            # Add a scene break if this is NOT the last file in the list.
            if i < len(file_paths) - 1:
                p = doc.add_paragraph('#')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def process_text_item(text_data, doc):
    """Processes a text item dictionary and writes its content to the document."""
    if 'file' in text_data:
        append_file_content(text_data['file'], doc)
    elif 'files' in text_data:
        file_paths = text_data.get('files', [])
        for i, file_path in enumerate(file_paths):
            append_file_content(file_path, doc)
            # Add a scene break if this is NOT the last file in the list.
            if i < len(file_paths) - 1:
                p = doc.add_paragraph('#')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def calculate_and_update_word_count(doc):
    """
    Calculates the total word count of the document and replaces the placeholder
    on the title page.
    """
    total_words = 0
    # Iterate through all paragraphs in the document to count words.
    for para in doc.paragraphs:
        # Simple word count logic: split by whitespace.
        total_words += len(para.text.split())

    # Round to the nearest 100 for the "Approx." count.
    rounded_words = int(round(total_words / 100.0)) * 100

    # Find the placeholder and replace it.
    for para in doc.tables[0].cell(0, 1).paragraphs:
        if "_WORD_COUNT_PLACEHOLDER_" in para.text:
            para.text = f"Approx. {rounded_words:,} words"

def compile_manuscript(config_file, output_dir):
    """Main function to read the YAML config and compile the manuscript into a .docx file."""
    print(f"Starting manuscript compilation from '{config_file}'...")

    try:
        with open(config_file, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)
    except FileNotFoundError:
        print(f"--> FATAL ERROR: The configuration file '{config_file}' was not found.")
        return
    except yaml.YAMLError as e:
        print(f"--> FATAL ERROR: There was an error parsing the YAML file: {e}")
        return

    if not os.path.exists(output_dir):
        print(f"Creating output directory: '{output_dir}'")
        os.makedirs(output_dir)
        
    # Get the output filename from the config, with a sensible default.
    metadata = config.get('metadata', {})
    story_type = metadata.get('story_type', 'novel').lower() # Default to 'novel' and ensure lowercase
    output_filename = metadata.get('file_name', 'manuscript.docx')
    output_path = os.path.join(output_dir, output_filename)

    # 1. Create Document and Set Styles
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Add the header to pages > 1
    setup_header(doc, config)

    # 2. Create the title page (if requested)
    include_title_page = metadata.get('include_title_page', True)
    if include_title_page:
        create_title_page(config, doc, story_type)

    is_first_content_item = True

    # 3. Process the main structure
    for item in config.get('structure', []):
        item_type = item.get('type')

        # Add a page break before any top-level part or chapter,
        # except for the very first one after the title page.
        if not is_first_content_item and story_type == 'novel':
            doc.add_page_break()
        elif not is_first_content_item and story_type == 'short_story':
            # For short stories, use line breaks instead of a page break
            doc.add_paragraph()
        is_first_content_item = False
        
        if item_type == 'part':
            part_title = smartypants.smartypants(item.get('title', 'Untitled Part').upper())
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(part_title)
            run.bold = True
            # Add space after part title
            doc.add_paragraph()
            
            # Process chapters within the part.
            chapters_in_part = item.get('content', [])
            for i, chapter in enumerate(chapters_in_part):
                # Add a page break ONLY if it's not the first chapter of the part.
                if i > 0 and story_type == 'novel':
                    doc.add_page_break()
                elif i > 0 and story_type == 'short_story':
                    doc.add_paragraph()
                process_chapter(chapter, doc, story_type)
        
        elif item_type == 'chapter':
            process_chapter(item, doc, story_type)
        
        elif item_type == 'text':
            process_text_item(item, doc)

        else:
            print(f"--> WARNING: Unknown structure type '{item_type}' found. Skipping.")
    
    # Add a final centered end mark on a new page.
    doc.add_page_break()
    p = doc.add_paragraph('#  #  #')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 4. Calculate final word count and update the title page
    if include_title_page:
        calculate_and_update_word_count(doc)

    # 5. Save the document
    doc.save(output_path)

    print(f"\nCompilation complete! Manuscript saved to '{output_path}'.")


def main():
    parser = argparse.ArgumentParser(
        description="Compiles a novel manuscript from a YAML configuration file into a single .docx file."
    )
    
    parser.add_argument(
        "config_file", 
        help="Path to the YAML configuration file (e.g., testbuild.yaml)."
    )
    
    parser.add_argument(
        "output_dir", 
        help="The directory where the compiled manuscript file will be saved (e.g., build/)."
    )
    
    args = parser.parse_args()
    compile_manuscript(args.config_file, args.output_dir)


if __name__ == "__main__":
    main()

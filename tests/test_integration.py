import os
import pytest
from docx import Document
from my_cli_tool.main import compile_manuscript

def test_compile_manuscript_integration(tmpdir):
    """
    An integration test that compiles a manuscript from a test config and
    verifies the output .docx file's structure and content.
    """
    # --- Setup ---
    output_dir = str(tmpdir)
    config_file = 'tests/test_config.yaml'

    # Expected output path based on the config file
    output_path = os.path.join(output_dir, 'test_output.docx')

    # --- Execution ---
    compile_manuscript(config_file, output_dir)

    # --- Verification ---
    assert os.path.exists(output_path), "The output docx file was not created."

    doc = Document(output_path)

    # 2. Verify Content
    # To make this robust, we'll collect all paragraph texts into a single string.
    full_text = "\\n".join([p.text for p in doc.paragraphs])

    # Check for title page content (in the main body)
    assert "MY TEST NOVEL" in full_text
    assert "by\nTest Author" in full_text

    # Check for word count (in the table)
    word_count_cell_text = doc.tables[0].cell(0, 1).text
    assert "Approx. 100 words" in word_count_cell_text

    # Check for chapter headings
    assert "Prologue" in full_text
    assert "Chapter 1: The First Chapter" in full_text

    # Check for content from the markdown files
    assert "This is the prologue." in full_text
    assert "first part of the first chapter" in full_text
    assert "second part of the first chapter" in full_text

    # Check for the scene break
    assert "#" in full_text

    # Check for the final end mark
    assert "#  #  #" in full_text

    # 3. Verify Formatting (Spot Check)
    # Check that the chapter 1 heading is centered
    chapter_1_heading_found = False
    for p in doc.paragraphs:
        if "Chapter 1: The First Chapter" in p.text:
            assert p.alignment.name == 'CENTER'
            chapter_1_heading_found = True
            break
    assert chapter_1_heading_found, "Chapter 1 heading was not found for format check."

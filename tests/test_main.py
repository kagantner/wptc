import os
import pytest
from docx import Document
from my_cli_tool.main import compile_manuscript

# Define the output directory for tests
TEST_OUTPUT_DIR = "test_build"

@pytest.fixture(scope="module", autouse=True)
def setup_teardown():
    """Create the test output directory before tests run and clean up after."""
    if not os.path.exists(TEST_OUTPUT_DIR):
        os.makedirs(TEST_OUTPUT_DIR)

    yield

    # Teardown: Clean up generated files
    for item in os.listdir(TEST_OUTPUT_DIR):
        if item.endswith(".docx"):
            os.remove(os.path.join(TEST_OUTPUT_DIR, item))
    os.rmdir(TEST_OUTPUT_DIR)

def docx_contains_text(path, text):
    """Helper function to check if a .docx file contains specific text."""
    doc = Document(path)
    for p in doc.paragraphs:
        if text in p.text:
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if text in cell.text:
                    return True
    return False

def test_manuscript_with_title_page():
    """
    Tests that a manuscript compiled with 'include_title_page: true'
    contains the author's legal name on the title page.
    """
    config_file = "testbuild_novel.yaml"
    compile_manuscript(config_file, TEST_OUTPUT_DIR)

    output_file = os.path.join(TEST_OUTPUT_DIR, "test_manuscript.docx")
    assert os.path.exists(output_file)

    # The legal name is a good proxy for the title page's existence
    assert docx_contains_text(output_file, "Jane Quinn Novelist")

def test_manuscript_without_title_page():
    """
    Tests that a manuscript compiled with 'include_title_page: false'
    does NOT contain the author's legal name from the title page.
    """
    config_file = "testbuild_short.yaml"
    compile_manuscript(config_file, TEST_OUTPUT_DIR)

    output_file = os.path.join(TEST_OUTPUT_DIR, "test_short_story.docx")
    assert os.path.exists(output_file)

    # The legal name should not be present if the title page is excluded
    assert not docx_contains_text(output_file, "Jane Quinn Novelist")

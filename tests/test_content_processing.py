import pytest
from unittest.mock import Mock, call, patch, mock_open
from my_cli_tool.main import (
    append_file_content,
    process_chapter,
    process_text_item,
)
from docx import Document

# --- Tests for append_file_content ---

@patch('builtins.open', new_callable=mock_open, read_data="Line 1\n\n# Heading\n% Comment\nLine 2")
@patch('my_cli_tool.main.add_smart_text')
def test_append_file_content_reads_and_processes_lines(mock_add_smart_text, mock_open):
    """
    Tests that append_file_content correctly reads a file, skips unwanted lines,
    and calls add_smart_text for valid content lines.
    """
    mock_doc = Mock()
    # When add_paragraph is called, it should return another mock
    # so that further attributes can be set on it (like paragraph_format)
    mock_paragraph = Mock()
    mock_doc.add_paragraph.return_value = mock_paragraph

    append_file_content("dummy/path.md", mock_doc)

    mock_open.assert_called_once_with("dummy/path.md", 'r', encoding='utf-8')

    # Check that add_smart_text was called with the correct, stripped lines
    calls = mock_add_smart_text.call_args_list
    assert len(calls) == 2
    assert calls[0].args[1] == "Line 1"
    assert calls[1].args[1] == "Line 2"

    # Check that new paragraphs were added to the doc
    assert mock_doc.add_paragraph.call_count == 2

@patch('builtins.open', side_effect=FileNotFoundError)
def test_append_file_content_handles_file_not_found(mock_open):
    """
    Tests that append_file_content gracefully handles a FileNotFoundError.
    """
    mock_doc = Mock()
    # We don't expect any calls to the document if the file doesn't exist
    append_file_content("nonexistent/file.md", mock_doc)
    mock_doc.add_paragraph.assert_not_called()

# --- Tests for process_chapter ---

@patch('my_cli_tool.main.append_file_content')
def test_process_chapter_with_number_and_title(mock_append):
    """Tests chapter processing with both a number and a title."""
    mock_doc = Mock()
    mock_paragraph = Mock()
    mock_doc.add_paragraph.return_value = mock_paragraph
    chapter_data = {'number': 1, 'title': 'The Beginning', 'file': 'chapter1.md'}

    process_chapter(chapter_data, mock_doc, 'novel')

    # Verify the heading was created correctly
    # The first call to add_paragraph is for the heading, the second for spacing
    assert mock_doc.add_paragraph.call_count == 2
    mock_doc.add_paragraph.assert_any_call() # For the spacing

    # Get the paragraph object the heading was added to
    heading_paragraph = mock_doc.add_paragraph.return_value
    heading_paragraph.add_run.assert_called_once_with("Chapter 1: The Beginning")

    mock_append.assert_called_once_with('chapter1.md', mock_doc)

@patch('my_cli_tool.main.append_file_content')
def test_process_chapter_with_multiple_files(mock_append):
    """Tests that chapters with multiple files have scene breaks added."""
    mock_doc = Mock()
    # Mock add_paragraph to return a mock paragraph object to check alignment
    mock_paragraph = Mock()
    mock_doc.add_paragraph.return_value = mock_paragraph

    chapter_data = {'files': ['scene1.md', 'scene2.md']}

    process_chapter(chapter_data, mock_doc, 'novel')

    # append_file_content should be called for each file
    assert mock_append.call_count == 2
    mock_append.assert_has_calls([call('scene1.md', mock_doc), call('scene2.md', mock_doc)])

    # A scene break ('#') should be added once, between the two files
    mock_doc.add_paragraph.assert_called_once_with('#')
    assert mock_paragraph.alignment is not None # Check that alignment was set

# --- Tests for process_text_item ---

@patch('my_cli_tool.main.append_file_content')
def test_process_text_item_single_file(mock_append):
    """Tests a simple text item with one file."""
    mock_doc = Mock(spec=Document)
    text_data = {'file': 'prologue.md'}

    process_text_item(text_data, mock_doc)

    mock_append.assert_called_once_with('prologue.md', mock_doc)

@patch('my_cli_tool.main.append_file_content')
def test_process_text_item_multiple_files(mock_append):
    """Tests that text items with multiple files have scene breaks."""
    mock_doc = Mock()
    mock_paragraph = Mock()
    mock_doc.add_paragraph.return_value = mock_paragraph

    text_data = {'files': ['intro1.md', 'intro2.md']}

    process_text_item(text_data, mock_doc)

    assert mock_append.call_count == 2
    mock_append.assert_has_calls([call('intro1.md', mock_doc), call('intro2.md', mock_doc)])

    # Check for the scene break
    mock_doc.add_paragraph.assert_called_once_with('#')
    assert mock_paragraph.alignment is not None
